from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Brand colour palette
BRAND_BLUE = RGBColor(0, 146, 209)
BRAND_NAVY = RGBColor(0, 51, 153)
BRAND_GOLD = RGBColor(255, 210, 86)
BRAND_DARK = RGBColor(26, 26, 46)
BRAND_GRAY = RGBColor(90, 90, 110)
BRAND_WHITE = RGBColor(255, 255, 255)
BRAND_LIGHT_GRAY = RGBColor(245, 247, 250)
BRAND_BORDER = RGBColor(208, 215, 222)

doc = Document()
sections = doc.sections[0]
sections.top_margin = Cm(2.0)
sections.bottom_margin = Cm(2.0)
sections.left_margin = Cm(2.5)
sections.right_margin = Cm(2.5)

# Set default font to Calibri
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = BRAND_DARK
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    fill = shading.find(qn('w:shd'))
    if fill is None:
        fill = OxmlElement('w:shd')
        shading.append(fill)
    fill.set(qn('w:fill'), color_hex)


def add_compact_heading(text, level):
    if level == 0:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = BRAND_NAVY
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        run2 = p2.add_run('━' * 50)
        run2.font.color.rgb = BRAND_GOLD
        run2.font.size = Pt(8)
        return p
    elif level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = BRAND_NAVY
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
        run2 = p2.add_run('━' * 30)
        run2.font.color.rgb = BRAND_BLUE
        run2.font.size = Pt(6)
        return p
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = BRAND_BLUE
        return p


add_compact_heading('Technical Assessment Submission: Data and AI Products (P2)', 0)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
run = p.add_run('Assessment: ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('#job2266-Technical Program Officer - Data and AI products for Humanitarian and Development Use-(P2)')
p.runs[1].font.size = Pt(10)

add_compact_heading('1. Live URLs', 1)
for label, url in [
    ('Public UI URL:', 'https://disaster-text-service-19133435506.us-east1.run.app/'),
    ('Public API /predict URL:', 'https://disaster-text-service-19133435506.us-east1.run.app/predict'),
    ('Public API /health URL:', 'https://disaster-text-service-19133435506.us-east1.run.app/health'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.color.rgb = BRAND_NAVY
    run.font.size = Pt(10)
    p.add_run(url).font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Additional API endpoints: ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('/predict_batch (batch classification), /stats (service metrics), /explain (model transparency), /sample (random tweet from dataset)').font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('GitHub Repository: ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('https://github.com/acandidate040/disaster-text-service').font.size = Pt(10)

add_compact_heading('2. Screenshot of Hosted UI', 1)
doc.add_paragraph('UI loaded in an incognito browser showing the browser address bar, sample input, prediction result, and confidence score:')
doc.add_picture('ui-screenshot.png', width=Inches(5.5))
doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

add_compact_heading('3. Quick API Sanity Check', 1)
p = doc.add_paragraph()
run = p.add_run('Command:')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
cmd = "curl -X POST https://disaster-text-service-19133435506.us-east1.run.app/predict \\\n  -H \"Content-Type: application/json\" \\\n  -d '{\"text\":\"Forest fire near La Ronge Sask. Canada\"}'"
p = doc.add_paragraph(cmd)
p.paragraph_format.left_indent = Inches(0.2)
p.paragraph_format.space_after = Pt(4)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_GRAY

p = doc.add_paragraph()
run = p.add_run('Response:')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p = doc.add_paragraph('{"label": 1, "score": 0.959883}')
p.paragraph_format.left_indent = Inches(0.2)
p.paragraph_format.space_after = Pt(6)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND_GRAY

add_compact_heading('4. Latency Test', 1)
doc.add_paragraph('A formal latency test was executed against the live /predict endpoint. 100 sequential POST requests were sent and the following statistical results were observed:')

latency_data = [
    ('Metric', 'Value'),
    ('Total requests', '100'),
    ('Successful', '100'),
    ('Errors', '0'),
    ('Mean', '0.386s'),
    ('p50 (median)', '0.328s'),
    ('p95', '0.759s'),
    ('p99', '1.035s'),
    ('Min', '0.275s'),
    ('Max', '1.035s'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'Value'
for cell in hdr:
    set_cell_shading(cell, '003399')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = BRAND_WHITE
            run.font.size = Pt(9)
        paragraph.paragraph_format.space_after = Pt(2)

for metric, value in latency_data[1:]:
    row = table.add_row().cells
    row[0].text = metric
    row[1].text = value
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
            paragraph.paragraph_format.space_after = Pt(2)

doc.add_paragraph('The p95 latency of 0.759s is well under the 2-second requirement, providing strong evidence that the service can reliably handle the grading workload.')

add_compact_heading('5. Score Distribution on Held-Out Data', 1)
doc.add_paragraph('The histogram below shows model confidence scores across a 500-example labelled validation sample that was not used to fit the classifier (250 disaster, 250 non-disaster). This provides a visual sanity check of class separation, although final grading will depend on the evaluator\'s held-out set.')
doc.add_picture('app/static/score_histogram.png', width=Inches(5.5))
doc.paragraphs[-1].paragraph_format.space_after = Pt(4)
doc.add_paragraph('Disaster tweets (red) cluster toward high confidence (mean: 0.769), while non-disaster tweets (green) cluster toward low confidence (mean: 0.214). The minimal overlap at the threshold confirms robust discriminative power.')

add_compact_heading('6. Anonymous Repo Link', 1)
p = doc.add_paragraph()
run = p.add_run('Repository: ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('https://github.com/acandidate040/disaster-text-service').font.size = Pt(10)

add_compact_heading('7. Technical Write-Up', 1)
add_compact_heading('What I built and why this approach', 2)
doc.add_paragraph('I built a single-container monolithic web service using Python (FastAPI). The system serves both the REST API endpoints and a static HTML/JavaScript frontend. For the model, I utilised a Scikit-Learn Pipeline featuring a FeatureUnion of word-level and character-level TF-IDF vectorizers, feeding into a Logistic Regression classifier with a finely tuned probability threshold.')
doc.add_paragraph('I chose this lightweight, classical ML architecture over a Deep Learning/Transformer approach to minimise cold-start and latency risk under the Cloud Run free-tier constraints (specifically min-instances=0 and p95 < 2s). Larger PyTorch or transformer-based containers can incur significant cold-start penalties, which increases the risk of timeout failures. This TF-IDF-based container keeps inference lightweight, reduces reliability risk, and exceeds the 0.70 F1 baseline used for grading.')

add_compact_heading('Model accuracy and validation', 2)
doc.add_paragraph('The model was trained on the Kaggle "Natural Language Processing with Disaster Tweets" training set using an 80/20 stratified split. On the validation split used for model selection, the classifier achieved an F1 score of 0.785 on the disaster class, exceeding the 0.70 baseline. The decision threshold was selected by grid search over validation probabilities to optimise F1.')

add_compact_heading('UI and hosting decisions', 2)
doc.add_paragraph("The application is deployed entirely on Google Cloud Run's free tier. Rather than hosting the frontend on a separate service, I served a vanilla HTML/CSS/JS frontend directly from FastAPI's static directory. This architectural choice eliminates CORS complexity, ensures the UI is strictly coupled to the exact API version deployed, and provides a highly resilient, single-command deployment pipeline. The UI features defensive JavaScript to handle empty inputs, loading states, and network errors gracefully. A random sample from the training dataset can be loaded via the /sample endpoint.")

add_compact_heading('AI tools used and what I validated or changed', 2)
doc.add_paragraph('I utilised a multi-agent AI workflow. ChatGPT 5.5 (High) and Gemini Pro 3.1 were used collaboratively to evaluate architectural trade-offs, manage the strict time constraints, and synthesise the safest deployment strategy (the Python Monolith). Devin and Kimi K2 were then used to scaffold the codebase and generate boilerplate logic for the Scikit-Learn pipeline and FastAPI routes. I independently verified the cross-validation logic, tuned the classification threshold to maximize the F1 score, and managed the Cloud Run deployment constraints manually. (See Appendix for full prompt log).')

add_compact_heading('Limitations and what I\'d improve with more time', 2)
doc.add_paragraph('Model Context & Sarcasm: Classical TF-IDF struggles with nuanced semantics, sarcasm, or highly localised slang. With more time, I would fine-tune a quantised foundational model (e.g., DistilBERT exported to ONNX) to improve contextual understanding.')
doc.add_paragraph('Calibration: The current confidence score represents the Logistic Regression probability, which is not perfectly calibrated to real-world likelihood. I would implement Platt scaling (via CalibratedClassifierCV) to make the UI confidence scores more accurate.')
doc.add_paragraph('Data Drift & MLOps: Disaster terminology evolves rapidly. I would integrate a feedback loop in the UI to flag false positives and set up a drift-monitoring dashboard.')

add_compact_heading('Service availability', 2)
doc.add_paragraph('The deployed Cloud Run service is configured to remain live and publicly accessible for at least 21 days following submission, ensuring ample time for grading and any follow-up evaluation. The container image and all source code are preserved in the public GitHub repository, allowing for rapid redeployment if necessary.')

add_compact_heading('8. References / Source Material', 1)
refs = [
    'Kaggle: Natural Language Processing with Disaster Tweets dataset and competition description.',
    'Google Cloud Run documentation: deploying containerized services and public unauthenticated services.',
    'FastAPI documentation: request validation and endpoint implementation.',
    'scikit-learn documentation: TfidfVectorizer, FeatureUnion, LogisticRegression, and model evaluation utilities.',
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(ref)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)

add_compact_heading('9. Appendix: AI Prompt Log', 1)

p = doc.add_paragraph()
run = p.add_run('AI Tools Disclosure Statement: ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('In accordance with the assessment instructions, AI agents were utilised to accelerate architectural planning, code generation, debugging, and documentation drafting. The final logic, system integration, deployment verification, and quality assurance were independently validated by the candidate.').font.size = Pt(10)

p = doc.add_paragraph()
p.add_run('Prompts from browser-based ChatGPT and Gemini sessions were transcribed manually from conversation history. The annex below represents all prompts that materially influenced architecture, code, or documentation.').font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('ChatGPT 5.5 (High) — browser-based chat: ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('Used for strategy synthesis, constraint evaluation (cold-start mitigation, latency analysis), and drafting project documentation, and final review.').font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Gemini Pro 3.1 — browser-based chat: ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('Used collaboratively alongside ChatGPT 5.5 for architectural planning, option comparison, documentation drafting and final review.').font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Devin IDE with Kimi K2.6 (free promotional model): ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('Used for directory scaffolding, generating the FastAPI boilerplate, writing the regex text-cleaning functions, constructing the static HTML frontend, deploying to Google Cloud Run, and creating the submission Word document.').font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('No paid AI or external LLM API is called by the deployed /predict endpoint. ')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run('The classifier runs locally inside the Cloud Run container.').font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run('Prompt Record:')
run.bold = True
run.font.color.rgb = BRAND_NAVY
run.font.size = Pt(10)
p.add_run(' (Iterative prompts used to formulate architecture and generate code).').font.size = Pt(10)

with open('../appendix-prompts.md', 'r') as f:
    md_content = f.read()

prompts = []
for line in md_content.split('\n'):
    match = re.match(r'\| ([^|]+) \| ([^|]+) \| (.+) \|', line.strip())
    if match and 'Date / Time' not in line and '---' not in line:
        date = match.group(1).strip()
        model = match.group(2).strip()
        text = match.group(3).strip()
        prompts.append((date, model, text))

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.LEFT

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Date / Time (EDT)'
hdr_cells[1].text = 'Model Used'
hdr_cells[2].text = 'Prompt Text'

for cell in hdr_cells:
    set_cell_shading(cell, '003399')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = BRAND_WHITE
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.space_before = Pt(2)

for date, model, text in prompts:
    row_cells = table.add_row().cells
    row_cells[0].text = date
    row_cells[1].text = model
    row_cells[2].text = text
    for i, cell in enumerate(row_cells):
        if i == 1:
            set_cell_shading(cell, 'f5f7fa')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = 'Calibri'
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.space_before = Pt(2)

table.columns[0].width = Inches(1.2)
table.columns[1].width = Inches(1.4)
table.columns[2].width = Inches(3.9)

add_compact_heading('10. Pre-submission Confirmation', 1)

checklist = [
    'Public UI URL opens in an incognito/private browser window with no login required.',
    'At least one successful prediction was generated through the UI.',
    'GET /health returns {"status":"ok"}.',
    'POST /predict returns JSON containing label as 0 or 1 and score as a float in [0, 1].',
    'Cloud Run service is configured with min-instances = 0.',
    'No GPU is used.',
    'No paid external LLM/API is called by /predict.',
    'Deployed service will remain live for at least 21 days.',
    'Public repository, commit history, README, and file metadata do not reveal identifying information.',
    'AI use has been disclosed, including the tool/model used, assisted parts, validation performed, and a prompt annex.',
]
for item in checklist:
    p = doc.add_paragraph()
    p.add_run('Confirmed: ' + item)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)

core_props = doc.core_properties
core_props.author = ''
core_props.title = ''
core_props.subject = ''
core_props.comments = ''
core_props.last_modified_by = ''
core_props.category = ''
core_props.keywords = ''

doc.save('Test Response.docx')
print('DOC_UPDATED')
